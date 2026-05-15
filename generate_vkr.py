# -*- coding: utf-8 -*-
"""
Конвертер markdown-текста ВКР в Word-документ (.docx)
с форматированием согласно всем требованиям оформления из prompt_docx.md.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# === ТЕКСТОВАЯ ОБРАБОТКА ===

PREPOSITIONS = [
    'в', 'на', 'с', 'к', 'у', 'о', 'и', 'а', 'но',
    'по', 'за', 'из', 'от', 'до', 'не', 'ни', 'об',
    'то', 'ко', 'во', 'со', 'же', 'бы', 'ли',
    'их', 'ей', 'ею', 'её', 'его'
]


def replace_dashes(text):
    """Замена длинных тире на средние."""
    return text.replace('\u2014', '\u2013')


def fix_hanging_prepositions(text):
    """Замена пробелов после предлогов на неразрывные."""
    nbsp = '\u00A0'
    for prep in PREPOSITIONS:
        # Замена после пробела или в начале строки
        pattern = r'(?<=\s)(' + re.escape(prep) + r')\s'
        text = re.sub(pattern, r'\1' + nbsp, text, flags=re.IGNORECASE)
        # В начале строки
        pattern = r'^(' + re.escape(prep) + r')\s'
        text = re.sub(pattern, r'\1' + nbsp, text, flags=re.IGNORECASE)
    return text


def process_text(text):
    """Обработка текста перед вставкой в документ."""
    text = replace_dashes(text)
    text = fix_hanging_prepositions(text)
    return text


# === НАСТРОЙКА ДОКУМЕНТА ===

def setup_document():
    """Создаёт документ с правильными полями и настройками."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # Стиль Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Убедимся что шрифт установлен для всех типов
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" '
            f'w:hAnsi="Times New Roman" w:cs="Times New Roman" '
            f'w:eastAsia="Times New Roman"/>'
        )
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    return doc


def add_page_numbers(doc):
    """Добавляет нумерацию страниц внизу по центру, 12пт."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = p.add_run()
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run2._element.append(instrText)

    run3 = p.add_run()
    run3.font.size = Pt(12)
    run3.font.name = 'Times New Roman'
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


# === ФУНКЦИИ ДОБАВЛЕНИЯ ЭЛЕМЕНТОВ ===

def add_paragraph(doc, text, align='justify', bold=False, font_size=14,
                  first_indent=1.25, space_before=0, space_after=0,
                  line_spacing=1.5):
    """Универсальная функция добавления параграфа."""
    p = doc.add_paragraph()

    alignments = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    p.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.first_line_indent = Cm(first_indent)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    if text:
        processed = process_text(text)
        run = p.add_run(processed)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'

    return p


def add_empty_line(doc, line_spacing=1.5):
    """Добавляет пустую строку (пустой параграф) с заданным интервалом."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_formatted_paragraph(doc, text, align='justify', first_indent=1.25,
                            font_size=14, line_spacing=1.5):
    """Добавляет параграф с обработкой inline-форматирования (жирный, курсив)."""
    p = doc.add_paragraph()

    alignments = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    p.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.first_line_indent = Cm(first_indent)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Разбираем inline форматирование: **жирный** и *курсив*
    # Паттерн: разбиваем текст на сегменты
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            content = process_text(content)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            content = part[1:-1]
            content = process_text(content)
            run = p.add_run(content)
            run.italic = True
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
        else:
            content = process_text(part)
            if content:
                run = p.add_run(content)
                run.font.size = Pt(font_size)
                run.font.name = 'Times New Roman'

    return p


# === ТАБЛИЦЫ ===

def add_table_borders(table):
    """Добавляет границы к таблице."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def create_table(doc, headers, rows):
    """Создаёт отформатированную таблицу."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(process_text(header))
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Данные
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Первый столбец — по левому краю, остальные — по центру
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(process_text(str(cell_text)))
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    return table


def parse_md_table(lines):
    """Парсит markdown-таблицу из списка строк. Возвращает (headers, rows)."""
    headers = []
    rows = []

    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Пропускаем разделительную строку (---)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)

    return headers, rows


# === ОСНОВНОЙ ПАРСЕР ===

def process_markdown(doc, md_text):
    """Обрабатывает markdown-текст и добавляет в документ."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    is_first_element = True  # Для первого элемента — не добавлять разрыв

    # Флаги для контекста
    last_was_chapter_heading = False
    last_was_paragraph_heading = False
    prev_paragraph_in_chapter = False  # Был ли уже хоть один параграф в текущей главе

    while i < len(lines):
        line = lines[i]

        # === Блоки кода (рисунки-схемы) ===
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Вставляем заглушку рисунка по центру
                add_paragraph(doc, '[Рисунок \u2013 см. оригинал]',
                              align='center', first_indent=0, font_size=12)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # === Горизонтальные разделители ===
        if line.strip() == '---':
            i += 1
            continue

        # === Пустые строки ===
        if not line.strip():
            i += 1
            continue

        # === Заголовок ВКР (# ВЫПУСКНАЯ...) ===
        if line.strip().startswith('# ВЫПУСКНАЯ'):
            text = re.sub(r'^#+\s*', '', line.strip())
            # Первая страница — титульная заглушка
            add_paragraph(doc, text.upper(), align='center',
                          bold=True, first_indent=0)
            is_first_element = False
            last_was_chapter_heading = False
            last_was_paragraph_heading = False
            i += 1
            continue

        # === Тема работы (## Тема:...) ===
        if line.strip().startswith('## Тема:'):
            text = re.sub(r'^##\s*', '', line.strip())
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            last_was_chapter_heading = False
            last_was_paragraph_heading = False
            i += 1
            continue

        # === ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ ===
        if re.match(r'^##\s*(ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ)', line.strip()):
            text = re.sub(r'^##\s*', '', line.strip())
            doc.add_page_break()
            add_paragraph(doc, text.upper(), align='center',
                          bold=True, first_indent=0)
            last_was_chapter_heading = True
            last_was_paragraph_heading = False
            prev_paragraph_in_chapter = False
            i += 1
            continue

        # === Заголовок главы (## ГЛАВА...) ===
        if line.strip().startswith('## ГЛАВА'):
            text = re.sub(r'^##\s*', '', line.strip())
            doc.add_page_break()
            add_paragraph(doc, text.upper(), align='center',
                          bold=True, first_indent=0)
            last_was_chapter_heading = True
            last_was_paragraph_heading = False
            prev_paragraph_in_chapter = False
            i += 1
            continue

        # === Вывод по главе (### Вывод по главе...) ===
        if line.strip().startswith('### Вывод по главе'):
            text = re.sub(r'^###\s*', '', line.strip())
            # Перед выводом — пустая строка
            add_empty_line(doc, line_spacing=1.5)
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            # После заголовка — текст начинается сразу (без пустой строки)
            last_was_chapter_heading = False
            last_was_paragraph_heading = True
            i += 1
            continue

        # === Заголовок параграфа (### X.X. ...) ===
        if re.match(r'^###\s+\d+\.\d+\.', line.strip()):
            text = re.sub(r'^###\s*', '', line.strip())

            if last_was_chapter_heading:
                # Между названием главы и первым параграфом: одинарный интервал
                add_empty_line(doc, line_spacing=1.0)
            else:
                # Перед параграфом (не первым в главе): пустая строка полуторная
                add_empty_line(doc, line_spacing=1.5)

            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            # После заголовка параграфа: НЕТ пустой строки
            last_was_chapter_heading = False
            last_was_paragraph_heading = True
            prev_paragraph_in_chapter = True
            i += 1
            continue

        # Сброс флагов после того как начался текст
        last_was_chapter_heading = False
        last_was_paragraph_heading = False

        # === Таблица markdown ===
        if line.strip().startswith('|'):
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                create_table(doc, headers, rows)
            continue

        # === Метка таблицы (Таблица X.X.X) ===
        if re.match(r'^Таблица\s+\d+', line.strip()):
            # Перед блоком таблицы — пустая строка ("таблица дышит")
            add_empty_line(doc, line_spacing=1.5)
            # По правому краю, НЕ жирный, без абзацного отступа
            add_paragraph(doc, line.strip(), align='right',
                          bold=False, first_indent=0, font_size=14)
            i += 1
            continue

        # === Название таблицы (жирный текст **...**) после метки ===
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            # Проверяем контекст: если предыдущая строка была "Таблица X.X.X"
            # то это название таблицы — по центру, НЕ жирный
            # Если же это подзаголовок в тексте — жирный по ширине
            # Эвристика: если текст длинный и нет рядом таблицы — это подзаголовок в тексте
            # Смотрим вперёд: если следующая строка — таблица markdown, это название таблицы
            next_non_empty = i + 1
            while next_non_empty < len(lines) and not lines[next_non_empty].strip():
                next_non_empty += 1

            if next_non_empty < len(lines) and lines[next_non_empty].strip().startswith('|'):
                # Это название таблицы — по центру, НЕ жирный (по требованиям)
                add_paragraph(doc, text, align='center',
                              bold=False, first_indent=0, font_size=14)
            else:
                # Это подзаголовок в тексте — жирный, по ширине, с абзацным отступом
                add_formatted_paragraph(doc, line.strip(), align='justify',
                                        first_indent=1.25, font_size=14)
            i += 1
            continue

        # === Источник таблицы/рисунка ===
        if line.strip().startswith('Источник:'):
            # Сразу под таблицей/рисунком, без пропуска строки, 12 пт
            add_paragraph(doc, line.strip(), align='justify',
                          font_size=12, first_indent=1.25)
            # После источника — пустая строка ("таблица/рисунок дышит")
            # Но проверим: если следующая строка — "Рис.", то не добавляем пустую
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx < len(lines) and re.match(r'^Рис\.\s+\d+', lines[next_idx].strip()):
                # Это блок рисунка — пустую строку добавим после "Рис."
                pass
            else:
                add_empty_line(doc, line_spacing=1.5)
            i += 1
            continue

        # === Подпись рисунка (Рис. X.X.X.) ===
        if re.match(r'^Рис\.\s+\d+', line.strip()):
            add_paragraph(doc, line.strip(), align='justify',
                          first_indent=1.25, font_size=14)
            # После подписи рисунка — пустая строка
            add_empty_line(doc, line_spacing=1.5)
            i += 1
            continue

        # === Элемент списка с тире ===
        if line.strip().startswith('- ') or line.strip().startswith('\u2013 '):
            text = re.sub(r'^[-\u2013]\s*', '', line.strip())
            # Убираем markdown-форматирование для обычного текста
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            add_paragraph(doc, '\u2013 ' + text, align='justify', first_indent=1.25)
            i += 1
            continue

        # === Нумерованный список X) ===
        if re.match(r'^\d+\)\s+', line.strip()):
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            add_paragraph(doc, text, align='justify', first_indent=1.25)
            i += 1
            continue

        # === Комментарий-заглушка *[...]*  ===
        if line.strip().startswith('*[') and line.strip().endswith(']*'):
            text = line.strip()[2:-2]
            add_paragraph(doc, '[' + text + ']', align='center',
                          first_indent=0, font_size=12)
            i += 1
            continue

        # === Обычный текст ===
        text = line.strip()

        # Проверяем наличие inline-форматирования
        if '**' in text or ('*' in text and not text.startswith('*[')):
            add_formatted_paragraph(doc, text, align='justify',
                                    first_indent=1.25, font_size=14)
        else:
            add_paragraph(doc, text, align='justify', first_indent=1.25)

        i += 1


def main():
    # Читаем markdown
    with open('/projects/sandbox/word-file/Текст_ВКР.md', 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Создаём документ
    doc = setup_document()
    add_page_numbers(doc)

    # Обрабатываем markdown и добавляем в документ
    process_markdown(doc, md_text)

    # Сохраняем
    output_path = '/projects/sandbox/word-file/ВКР.docx'
    doc.save(output_path)
    print(f"Документ успешно создан: {output_path}")


if __name__ == '__main__':
    main()
