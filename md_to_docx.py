# -*- coding: utf-8 -*-
"""
Конвертер markdown-текста ВКР в Word-документ (.docx) 
с форматированием согласно требованиям оформления.
"""

import re
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def setup_document():
    """Создаёт документ с правильными полями и настройками."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Стиль Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Шрифт
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" '
            f'w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        )
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
    
    return doc


def add_page_numbers(doc):
    """Добавляет нумерацию страниц внизу по центру, 12пт."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run()
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    
    run2 = p.add_run()
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run2._element.append(instrText)
    
    run3 = p.add_run()
    run3.font.size = Pt(12)
    run3.font.name = 'Times New Roman'
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def add_paragraph(doc, text, align='justify', bold=False, font_size=14,
                  first_indent=1.25, space_before=0, space_after=0):
    """Универсальная функция добавления параграфа."""
    p = doc.add_paragraph()
    
    alignments = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }
    p.alignment = alignments.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.first_line_indent = Cm(first_indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
    
    return p


def add_table_borders(table):
    """Добавляет границы к таблице."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def create_table(doc, headers, rows):
    """Создаёт отформатированную таблицу."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    # Данные
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    
    return table


def parse_md_table(lines):
    """Парсит markdown-таблицу из списка строк. Возвращает (headers, rows)."""
    headers = []
    rows = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Пропускаем разделительную строку (---)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    
    return headers, rows


def process_markdown(doc, md_text):
    """Обрабатывает markdown-текст и добавляет в документ."""
    lines = md_text.split('\n')
    i = 0
    is_first_chapter = True
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Обработка блоков кода (пропускаем - это схемы/рисунки-заглушки)
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                # Добавляем заглушку для рисунка
                add_paragraph(doc, '[Рисунок - см. оригинал]', 
                            align='center', first_indent=0, font_size=12)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Пропускаем горизонтальные разделители
        if line.strip() == '---':
            i += 1
            continue
        
        # Пропускаем пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Заголовок главы (## ГЛАВА...)
        if line.strip().startswith('## ГЛАВА') or line.strip().startswith('# ВЫПУСКНАЯ'):
            text = re.sub(r'^#+\s*', '', line.strip())
            if is_first_chapter and line.strip().startswith('# '):
                # Первый заголовок - без разрыва страницы
                doc.add_page_break()
                is_first_chapter = False
            else:
                doc.add_page_break()
            
            p = add_paragraph(doc, text.upper(), align='center', 
                            bold=True, first_indent=0)
            i += 1
            continue
        
        # Тема работы (## Тема:...)
        if line.strip().startswith('## Тема:'):
            text = re.sub(r'^##\s*', '', line.strip())
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            i += 1
            continue
        
        # Заголовок ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и пр.
        if line.strip().startswith('## ВВЕДЕНИЕ') or line.strip().startswith('## ЗАКЛЮЧЕНИЕ'):
            text = re.sub(r'^##\s*', '', line.strip())
            doc.add_page_break()
            add_paragraph(doc, text.upper(), align='center', bold=True, first_indent=0)
            i += 1
            continue
        
        # Вывод по главе
        if line.strip().startswith('### Вывод по главе'):
            text = re.sub(r'^###\s*', '', line.strip())
            # Пустая строка перед
            add_paragraph(doc, '', first_indent=0)
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            # Пустая строка после
            add_paragraph(doc, '', first_indent=0)
            i += 1
            continue
        
        # Заголовок параграфа (### 1.1., ### 1.2., etc.)
        if re.match(r'^###\s+\d+\.\d+\.', line.strip()):
            text = re.sub(r'^###\s*', '', line.strip())
            # Один интервал между главой и параграфом (просто новая строка)
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            # Пустая строка после заголовка параграфа
            add_paragraph(doc, '', first_indent=0)
            i += 1
            continue
        
        # Таблица markdown
        if line.strip().startswith('|'):
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(table_lines)
            if headers and rows:
                create_table(doc, headers, rows)
                # Пустая строка после таблицы
                add_paragraph(doc, '', first_indent=0)
            continue
        
        # Метка таблицы (Таблица X.X.X)
        if re.match(r'^Таблица\s+\d+', line.strip()):
            add_paragraph(doc, line.strip(), align='right', first_indent=0)
            i += 1
            continue
        
        # Источник таблицы/рисунка
        if line.strip().startswith('Источник:'):
            add_paragraph(doc, line.strip(), align='justify', 
                        font_size=12, first_indent=1.25)
            i += 1
            continue
        
        # Подпись рисунка (Рис. X.X.X.)
        if re.match(r'^Рис\.\s+\d+', line.strip()):
            add_paragraph(doc, line.strip(), align='justify', first_indent=1.25)
            i += 1
            continue
        
        # Жирный текст (подзаголовок внутри параграфа **текст**)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            # Это название таблицы по центру
            add_paragraph(doc, text, align='center', bold=True, first_indent=0)
            i += 1
            continue
        
        # Элемент списка с тире
        if line.strip().startswith('- ') or line.strip().startswith('\u2013 '):
            text = re.sub(r'^[-–]\s*', '', line.strip())
            # Убираем markdown-жирность
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_paragraph(doc, '\u2013 ' + text, align='justify', first_indent=1.25)
            i += 1
            continue
        
        # Нумерованный список
        if re.match(r'^\d+\)\s+', line.strip()):
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            add_paragraph(doc, text, align='justify', first_indent=1.25)
            i += 1
            continue
        
        # Комментарий в квадратных скобках (заглушка)
        if line.strip().startswith('*[') and line.strip().endswith(']*'):
            text = line.strip().strip('*[]')
            add_paragraph(doc, '[' + text + ']', align='center', 
                        first_indent=0, font_size=12)
            i += 1
            continue
        
        # Обычный текст
        text = line.strip()
        # Убираем markdown-форматирование
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        
        if text:
            add_paragraph(doc, text, align='justify', first_indent=1.25)
        
        i += 1


def main():
    # Читаем markdown
    with open('/projects/sandbox/word-file/Текст_ВКР.md', 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # Создаём документ
    doc = setup_document()
    add_page_numbers(doc)
    
    # Обрабатываем markdown и добавляем в документ
    process_markdown(doc, md_text)
    
    # Сохраняем
    output_path = '/projects/sandbox/word-file/ВКР.docx'
    doc.save(output_path)
    print(f"Документ успешно создан: {output_path}")


if __name__ == '__main__':
    main()
